import os
import re
import smtplib
from email.message import EmailMessage
from pathlib import Path

import requests
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1:8b")


def extract_text_from_file(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = ""
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".docx":
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    if suffix == ".csv":
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return f"[CSV file detected but could not be read: {path.name}]"

    if suffix == ".xlsx":
        return f"[Structured Excel file detected: {path.name}]"

    return f"[Unsupported file type: {path.name}]"


def normalize_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text


def extract_metrics(text: str) -> dict:
    metrics = {}

    patterns = {
        "total_responses": r"Total Check-in Responses:\s*(\d+)|Total Responses.*?(\d+)",
        "actively_working": r"Actively Working\s+(\d+)|Actively Working.*?(\d+)",
        "looking_to_contribute": r"Looking to Contribute\s+(\d+)|Looking to Contribute.*?(\d+)",
        "stepped_away": r"Stepped Away\s+(\d+)|Stepped Away.*?(\d+)",
        "need_1to1_support": r"Need 1:1 Support\s+(\d+)|Need 1:1 Support.*?(\d+)",
        "have_minor_blockers": r"Have Minor Blockers\s+(\d+)|Have Minor Blockers.*?(\d+)",
        "waiting_for_direction": r"Waiting for Direction\s+(\d+)|Waiting for Direction.*?(\d+)",
        "week1_response_rate": r"Week 1.*?(\d+%)",
        "week2_response_rate": r"Week 2.*?(\d+%)",
        "week1_nonrespondents": r"Week 1.*?Non-Respondents.*?(\d+)|Week 1.*?did not respond.*?(\d+)",
        "week2_nonrespondents": r"Week 2.*?Non-Respondents.*?~?(\d+)|Approximately\s*~?(\d+)\s*volunteers did not respond",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL | re.MULTILINE)
        if match:
            value = next((g for g in match.groups() if g), None)
            if value:
                metrics[key] = value

    return metrics


def extract_relevant_sections(text: str) -> dict:
    text = normalize_text(text)
    sections = {}

    section_patterns = {
        "executive_summary": r"1\.\s*Executive Summary(.*?)(?=2\.\s*Volunteer Roster|$)",
        "volunteer_roster": r"2\.\s*Volunteer Roster.*?(?=3\.\s*Week 1 Detailed Report|$)",
        "week1": r"3\.\s*Week 1 Detailed Report.*?(?=4\.\s*Week 2 Detailed Report|$)",
        "week2": r"4\.\s*Week 2 Detailed Report.*?(?=5\.\s*Team Progress by Function|$)",
        "team_progress": r"5\.\s*Team Progress by Function(.*?)(?=6\.\s*Support Needs & Action Items|$)",
        "support_needs": r"6\.\s*Support Needs & Action Items(.*?)(?=7\.\s*Leadership Updates & Decisions|$)",
        "leadership_updates": r"7\.\s*Leadership Updates & Decisions(.*?)(?=8\.\s*Overview & Next Week Action Plan|$)",
        "overview_action_plan": r"8\.\s*Overview & Next Week Action Plan(.*)$",
    }

    for key, pattern in section_patterns.items():
        match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
        sections[key] = match.group(1).strip() if match else ""

    if not any(sections.values()):
        sections["fallback"] = text[:9000]

    sections["metrics"] = extract_metrics(text)
    return sections


def build_structured_context(file_paths: list[str]) -> str:
    all_context = []

    for file_path in file_paths:
        extracted = extract_text_from_file(file_path)
        sections = extract_relevant_sections(extracted)

        file_block = [f"FILE: {Path(file_path).name}"]

        metrics = sections.get("metrics", {})
        if metrics:
            file_block.append("METRICS:")
            for k, v in metrics.items():
                file_block.append(f"- {k}: {v}")

        for section_name in [
            "executive_summary",
            "team_progress",
            "support_needs",
            "leadership_updates",
            "overview_action_plan",
            "volunteer_roster",
            "week1",
            "week2",
            "fallback",
        ]:
            content = sections.get(section_name, "")
            if content:
                pretty_name = section_name.replace("_", " ").title()
                file_block.append(f"{pretty_name}:")
                file_block.append(content[:5000])

        all_context.append("\n".join(file_block))

    return "\n\n".join(all_context)


def clean_report_text(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("### ", "")
    text = text.replace("## ", "")
    text = text.replace("# ", "")
    text = text.replace("* ", "")
    text = text.replace("This appears to be", "")
    text = text.replace("I'll break down the key points into sections:", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def get_prompt_template(report_type: str, month: str, notes: str, structured_context: str) -> str:
    common_rules = f"""
Use only facts present in the structured source.
Do not invent names, owners, or project details.
If detailed operational information is present, summarize it at the correct level for the selected report type.

User Notes:
{notes if notes.strip() else "No additional notes provided."}

Structured Source:
{structured_context[:22000]}
"""

    if report_type == "executive":
        return f"""
You are writing a founder-level executive report for Guardians Embrace.

Return ONLY the final report.
Do not add commentary.
Do not explain what the document is.
Do not use markdown or asterisks.

Use EXACTLY this structure:

Executive Summary:
Write 3 concise sentences covering overall momentum, top concern, and leadership priority.

Volunteer Engagement:
1. ...
2. ...
3. ...
4. ...

Project Status:
1. ...
2. ...
3. ...
4. ...

Critical Issues:
1. ...
2. ...
3. ...
4. ...

Next Steps:
1. ...
2. ...
3. ...
4. ...

Key Takeaways:
Write 3 concise sentences focused on founder-level conclusions.

STRICT EXECUTIVE FILTER RULES:
- This is a leadership-level report, not an operational report.
- Do NOT include volunteer rosters.
- Do NOT list individual contributor names in Volunteer Engagement.
- Do NOT include contributor names in Project Status unless they are leadership owners or decision-makers.
- Do NOT describe person-level activities, meeting histories, or task-by-task narratives.
- Prefer counts, trends, risks, ownership gaps, project status, and leadership actions.
- Summarize work at the team or project level, not at the person-task level.
- Do NOT include unnecessary weekly detail if it can be aggregated into one executive insight.
- If the source contains detailed names and activities, compress them into higher-level takeaways.

SECTION RULES:
- Executive Summary must focus on overall momentum, biggest concern, and immediate leadership attention.
- Volunteer Engagement must include only counts, participation trends, support needs, direction gaps, blockers, and non-response issues.
- Project Status must summarize projects by status such as On Track, In Progress, Early Stage, Pending Alignment, or Blocked.
- Project Status must mirror the source's project-level framing when available.
- Prefer project-level summaries such as Geo-Mapping & Church Data, Donor Management & Dashboards, IT Infrastructure & Tools, Portal Development, Executive Dashboard & Reporting, and Donor Communication Project.
- Critical Issues must focus on risks, stalled alignment, unclear ownership, support needs, communication gaps, and unresolved volunteer status.
- Next Steps must focus on leadership actions, follow-ups, and high-priority coordination steps.
- Key Takeaways must sound concise, strategic, and founder-ready.

SOURCE FIDELITY RULES:
- Use only facts explicitly supported by the structured source.
- Mention exact numbers where available.
- If response-rate decline is present, state it clearly.
- If non-responding volunteers are present, state it clearly.
- If donor communication is pending, stalled, or awaiting alignment, state it clearly.
- If dashboard ownership is unclear, state it clearly.
- If support needs or direction gaps are present, state them clearly.
- If the source includes donor communication alignment issues, dashboard ownership ambiguity, or pending project clarification, those must appear in Project Status or Critical Issues.
- Do NOT say "Not specified" when the source provides enough evidence to describe an executive issue.
- If detailed source text shows ownership ambiguity, summarize it as an ownership or alignment issue instead of copying task-level detail.

IMPORTANT EXECUTIVE WRITING BEHAVIOR:
- Prefer: "Donor Communication remains pending alignment and requires ownership clarity."
- Prefer: "Executive Dashboard & Reporting is in early stage and depends on cross-team data gathering."
- Prefer: "Volunteer engagement needs attention due to declining response rate and unresolved non-respondent status."
- Prefer: "Three volunteers require 1:1 support and four are still waiting for direction."
- Prefer: "Portal Development is in progress with a minor blocker related to Google Drive integration."
- Avoid: long lists of names, task histories, raw roster data, or detailed contributor attribution.
- Avoid: repeating the same fact across all sections unless necessary.

{common_rules}
"""

    if report_type == "operational":
        return f"""
You are writing an internal operational report for Guardians Embrace.

Return ONLY the final report.
Do not add commentary.
Do not use markdown or asterisks.

Use EXACTLY this structure:

Operational Summary:
Write 3 concise sentences.

Team Activity:
1. ...
2. ...
3. ...
4. ...

Project Progress:
1. ...
2. ...
3. ...
4. ...

Blockers:
1. ...
2. ...
3. ...
4. ...

Immediate Follow-Up:
1. ...
2. ...
3. ...
4. ...

Operational Takeaways:
Write 3 concise sentences.

Operational rules:
- Contributor names may be included if they are clearly supported by the source
- Include practical execution details
- Highlight waiting-for-direction issues, access/context gaps, and blockers

{common_rules}
"""

    if report_type == "action_plan":
        return f"""
You are writing a concise action plan for Guardians Embrace.

Return ONLY the final report.
Do not add commentary.
Do not use markdown or asterisks.

Use EXACTLY this structure:

Priority Actions:
1. ...
2. ...
3. ...
4. ...

Blockers:
1. ...
2. ...
3. ...
4. ...

Ownership Needed:
1. ...
2. ...
3. ...
4. ...

Immediate Next Moves:
1. ...
2. ...
3. ...
4. ...

Rules:
- Be direct and actionable
- Focus only on what needs to be done
- Avoid descriptive summaries unless needed for clarity

{common_rules}
"""

    if report_type == "speechify":
        return f"""
You are writing a spoken executive update for Guardians Embrace.

Return ONLY the final script.
Do not add commentary.
Do not use markdown or asterisks.

=====================
FORMAT RULES
=====================

- Start with: Hello,
- DO NOT say "Good morning", "Good evening", or "everyone"
- DO NOT include "Best regards", signatures, or names
- DO NOT include placeholders like [Your Name]

=====================
CONTENT RULES
=====================

- This is a spoken report, not an email
- Keep it natural and smooth when read aloud
- Use short, clear sentences
- Use simple transitions between ideas

- DO NOT include contributor names
- DO NOT include technical implementation details
- DO NOT include task-level descriptions

- Focus on:
  → overall progress
  → key updates
  → concerns
  → priorities

=====================
MANDATORY SIGNALS
=====================

Include if present:
- Response rate decline
- Non-responding volunteers
- Support needs
- Direction gaps
- Project alignment issues

=====================
STRUCTURE (NATURAL FLOW)
=====================

1. Greeting (Hello)
2. Overall progress
3. Key project updates
4. Main concerns
5. Immediate priorities
6. Forward-looking statement



{common_rules}
"""

    if report_type == "dashboard_summary":
        return f"""
You are writing a dashboard-style summary for Guardians Embrace.

Return ONLY the final report.
Do not add commentary.
Do not use markdown or asterisks.

Use EXACTLY this structure:

Overview Card:
Write 2 concise sentences.

Volunteer Card:
1. ...
2. ...
3. ...

Projects Card:
1. ...
2. ...
3. ...

Risks Card:
1. ...
2. ...
3. ...

Next Actions Card:
1. ...
2. ...
3. ...

Rules:
- Very concise
- Card-style summary
- Use metrics and executive signals
- Do not include names
- Make it suitable for linking to deeper drill-downs later

{common_rules}
"""

    return f"""
You are writing a concise report for Guardians Embrace.

Return ONLY the final report.
Do not add commentary.
Do not use markdown.

Write a structured report with clear sections and concise content.

{common_rules}
"""


def generate_report_with_ollama(report_type: str, month: str, notes: str, structured_context: str) -> str:
    prompt = get_prompt_template(report_type, month, notes, structured_context)

    response = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json={
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0.0,
                "top_p": 0.9,
                "num_predict": 1500
            }
        },
        timeout=300
    )

    response.raise_for_status()
    data = response.json()

    if "response" not in data:
        raise ValueError(f"Unexpected Ollama response: {data}")

    return clean_report_text(data["response"])


def save_word_doc(month: str, report_type: str, report_text: str) -> str:
    doc = Document()
    doc.add_heading(
        f"Guardians Embrace {report_type.replace('_', ' ').title()} - {month}",
        0,
    )

    known_headings = {
        "executive summary:",
        "volunteer engagement:",
        "project status:",
        "critical issues:",
        "next steps:",
        "key takeaways:",
        "operational summary:",
        "team activity:",
        "project progress:",
        "blockers:",
        "immediate follow-up:",
        "operational takeaways:",
        "priority actions:",
        "ownership needed:",
        "immediate next moves:",
        "overview card:",
        "volunteer card:",
        "projects card:",
        "risks card:",
        "next actions card:",
        "opening",
        "main progress",
        "main concerns",
        "immediate priorities",
        "closing",
    }

    for line in report_text.split("\n"):
        clean = line.strip()
        if not clean:
            continue

        if clean.lower() in known_headings:
            doc.add_heading(clean.rstrip(":"), level=1)
        else:
            doc.add_paragraph(clean)

    file_name = f"GE_{report_type.replace(' ', '_').title()}_{month.replace(' ', '_')}.docx"
    output_path = OUTPUT_DIR / file_name
    doc.save(output_path)
    return str(output_path)


def send_email_with_attachment(to_email: str, subject: str, body: str, attachment_path: str):
    smtp_server = os.getenv("SMTP_SERVER")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_email = os.getenv("SMTP_EMAIL")
    smtp_password = os.getenv("SMTP_PASSWORD")

    if not smtp_server or not smtp_email or not smtp_password:
        raise ValueError("SMTP configuration missing in .env")

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp_email
    msg["To"] = to_email
    msg.set_content(body)

    with open(attachment_path, "rb") as f:
        file_data = f.read()
        file_name = Path(attachment_path).name

    msg.add_attachment(
        file_data,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name,
    )

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_email, smtp_password)
        server.send_message(msg)


def process_files_and_generate_report(
    month: str,
    email: str,
    notes: str,
    report_type: str,
    file_paths: list[str],
) -> str:
    structured_context = build_structured_context(file_paths)

    report_text = generate_report_with_ollama(
        report_type=report_type,
        month=month,
        notes=notes,
        structured_context=structured_context,
    )

    word_file = save_word_doc(
        month=month,
        report_type=report_type,
        report_text=report_text,
    )

    send_email_with_attachment(
        to_email=email,
        subject=f"Guardians Embrace {report_type.replace('_', ' ').title()} - {month}",
        body=f"Attached is the generated {report_type.replace('_', ' ')} for {month}.",
        attachment_path=word_file,
    )

    return f"SUCCESS: Report generated and emailed -> {word_file}"